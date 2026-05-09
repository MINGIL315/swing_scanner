"""KR 시장 분석 모듈 흐름 Word 문서 생성기."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_FONT_BODY = "맑은 고딕"
_FONT_CODE = "Consolas"


def _set_run_font(run, name: str = _FONT_BODY, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def _add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    _set_run_font(run, _FONT_BODY, 16 if level == 1 else 13 if level == 2 else 11.5, bold=True)


def _add_para(doc, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, _FONT_BODY, 10.5, bold=bold)
    run.italic = italic


def _add_code(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    _set_run_font(run, _FONT_CODE, 9.5)


def _add_table(doc, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_run_font(run, _FONT_BODY, 10.0, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        cells = tbl.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(val)
            _set_run_font(
                run,
                _FONT_CODE if (c_idx == 0 and val and ("/" in val or val.endswith(".py"))) else _FONT_BODY,
                9.5,
            )
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build() -> Path:
    doc = Document()

    # ── 표지 ──────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Swing Scanner")
    _set_run_font(run, _FONT_BODY, 22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("KR 시장 분석 모듈 — 전체 흐름 가이드")
    _set_run_font(run, _FONT_BODY, 14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("작성일: 2026-05-09  |  대상 브랜치: feat/web-entry-signal-and-4h-toggle")
    _set_run_font(run, _FONT_BODY, 10)
    run.italic = True

    doc.add_paragraph()

    # ── 1. 개요 ────────────────────────────────────────────────────
    _add_heading(doc, "1. 개요", 1)
    _add_para(doc, "Swing Scanner 는 KOSPI200 약 200개 종목을 매일 자동 스캔해 4가지 상승 패턴 (쌍바닥·골든크로스·박스권 돌파·눌림목) 을 탐지하고, 신뢰도 점수와 진입가/손절가/목표가를 제시하는 스윙매매 차트 발굴 시스템이다.")
    _add_para(doc, "본 문서는 한국 시장 (KR) 분석 모듈에 한정해, 데이터 수집부터 리포트 출력까지의 전체 흐름과 각 단계에서 사용되는 모듈 파일 경로를 정리한다.", italic=True)

    _add_heading(doc, "1.1 멀티 타임프레임 철학", 2)
    _add_table(
        doc,
        ["프레임", "역할", "판단 기준"],
        [
            ["주봉",                "큰 추세 확인",        "최근 12주 종가 기울기 + 20주선 위/아래"],
            ["일봉",                "패턴 탐지",          "4가지 패턴 정의 (쌍바닥/골든크로스/박스권 돌파/눌림목)"],
            ["4시간봉 (60분 × 4)",  "진입 타이밍 평가",    "거래량·캔들·RSI 다이버전스"],
        ],
        col_widths_cm=[3.5, 4.0, 9.0],
    )
    _add_para(doc, "→ 위에서 아래로 통과한 종목만 후보. 주봉이 하락 추세면 일봉 패턴이 보여도 버린다 (가짜 패턴 노이즈 제거).")

    # ── 2. 폴더 구조 ──────────────────────────────────────────────
    _add_heading(doc, "2. KR 모듈 폴더 구조", 1)
    _add_code(doc,
        "scanner/\n"
        "├── cli.py                       # typer 기반 CLI 진입점 (공통)\n"
        "├── config.py                    # 임계값 상수 + 환경 설정 (공통)\n"
        "├── pipeline.py                  # 일일 스캔 오케스트레이터 (공통)\n"
        "├── data_pipeline.py             # 통합 데이터 수집 dispatcher (KR/US 분기)\n"
        "│\n"
        "├── kr/                          # 한국 시장 자체 완결 분석 모듈\n"
        "│   ├── fetcher.py               # KIS OpenAPI — 일봉/1분봉/재무\n"
        "│   ├── universe.py              # KOSPI200 구성종목 갱신\n"
        "│   ├── scanner.py               # 핵심 분석 — analyze_ticker / scan_universe\n"
        "│   ├── intraday.py              # 1분봉 → 60min/4h 합성\n"
        "│   ├── quote_align.py           # 가격대별 호가단위 정합\n"
        "│   ├── indicators/\n"
        "│   │   ├── moving_average.py    # SMA / EMA\n"
        "│   │   ├── rsi.py               # RSI(14)\n"
        "│   │   ├── macd.py              # MACD(12,26,9)\n"
        "│   │   └── volume.py            # volume_ratio (당일/평균)\n"
        "│   ├── patterns/\n"
        "│   │   ├── base.py              # PatternDetector 추상 + EntrySignal/PatternResult\n"
        "│   │   ├── trend.py             # 주봉 추세 판정 (uptrend/downtrend/sideways)\n"
        "│   │   ├── double_bottom.py     # 쌍바닥\n"
        "│   │   ├── golden_cross.py      # 골든크로스\n"
        "│   │   ├── box_breakout.py      # 박스권 돌파\n"
        "│   │   └── pullback.py          # 눌림목\n"
        "│   ├── scoring/\n"
        "│   │   └── scorer.py            # 신뢰도 점수 (주봉30+패턴25+거래량20+정배열15+RSI10)\n"
        "│   ├── filtering/\n"
        "│   │   ├── volume_filter.py     # 거래량 모멘텀 (5일평균 > 20일평균)\n"
        "│   │   └── fundamental_filter.py # 시가총액 1000억원 이상\n"
        "│   ├── backtest/\n"
        "│   │   └── engine.py            # 과거 ScanResult 재생 시뮬레이션\n"
        "│   └── reports/\n"
        "│       ├── html_report.py       # 정적 HTML (Jinja2 + LightweightCharts)\n"
        "│       ├── excel_export.py      # CSV/XLSX 내보내기\n"
        "│       └── comment_generator.py # 패턴별 한국어 코멘트\n"
        "│\n"
        "├── db/                          # 공통 — SQLAlchemy 모델·세션\n"
        "│   ├── models.py                # Universe / OHLCVDaily / OHLCVWeekly /\n"
        "│   │                            #   OHLCVIntraday / ScanResult / Fundamental\n"
        "│   ├── session.py               # get_session() 컨텍스트 매니저\n"
        "│   ├── repository.py            # save_scan_results / get_scan_results\n"
        "│   ├── migrations.py            # init_database (SQLite WAL)\n"
        "│   └── universe_db.py           # get_active_tickers\n"
        "│\n"
        "└── api/                         # 공통 — FastAPI\n"
        "    ├── main.py                  # FastAPI app\n"
        "    └── routers/\n"
        "        └── stocks.py            # /api/stocks/{ticker}/{ohlcv,analysis}\n"
    )

    # ── 3. 전체 흐름 시퀀스 ────────────────────────────────────────
    _add_heading(doc, "3. 일일 스캔 전체 흐름 시퀀스", 1)
    _add_para(doc, "CLI 명령어 `scanner scan --market kr` 실행 시의 단계별 호출 순서다. 각 단계가 어느 모듈을 호출하는지 명시한다.", italic=True)
    _add_code(doc,
        "$ scanner scan --market kr\n"
        "    │\n"
        "    ├─[1] 유니버스 갱신 (7일 미갱신 시)        → scanner/kr/universe.py:update_kospi200\n"
        "    │       └─ KIS 종목마스터 zip 다운로드 + 활성 종목 DB upsert\n"
        "    │\n"
        "    ├─[2] 데이터 수집 (--skip-fetch 미사용 시) → scanner/data_pipeline.py:run_data_pipeline\n"
        "    │       ├─ 일봉/재무 fetch                  → scanner/kr/fetcher.py\n"
        "    │       └─ 1분봉 fetch (--with-intraday)    → scanner/kr/fetcher.py:fetch_minute_ohlcv\n"
        "    │       │   (DB: ohlcv_daily / ohlcv_intraday / fundamentals 테이블)\n"
        "    │\n"
        "    ├─[3] 일봉 + 분봉 + 재무 로드             → scanner/pipeline.py\n"
        "    │       ├─ _load_daily_dfs()                : 500일 일봉\n"
        "    │       ├─ _load_intraday_dfs()             : 30일 1분봉 → resample_to_4h(drop_partial=True)\n"
        "    │       │       └─ scanner/kr/intraday.py:resample_to_4h\n"
        "    │       │             (origin='start_day', offset='9h' → 09:00/13:00 그룹)\n"
        "    │       └─ _load_fundamentals() : 시가총액 + 재무\n"
        "    │\n"
        "    ├─[4] 패턴 분석 (병렬, ThreadPoolExecutor) → scanner/kr/scanner.py:analyze_ticker\n"
        "    │       ├─ 주봉 추세 판정                    → scanner/kr/patterns/trend.py:detect_weekly_trend\n"
        "    │       ├─ 4가지 패턴 detect()              → scanner/kr/patterns/{double_bottom,golden_cross,\n"
        "    │       │                                                          box_breakout,pullback}.py\n"
        "    │       ├─ 신뢰도 점수 계산                  → scanner/kr/scoring/scorer.py:calculate_confidence_score\n"
        "    │       ├─ 진입 신호 평가 (4시간봉)         → detector.entry_signal(daily_df, intraday_df=4h봉)\n"
        "    │       ├─ 거래량 필터                      → scanner/kr/filtering/volume_filter.py\n"
        "    │       └─ 재무 필터 (옵션)                 → scanner/kr/filtering/fundamental_filter.py\n"
        "    │\n"
        "    ├─[5] DB 저장                             → scanner/db/repository.py:save_scan_results\n"
        "    │       (테이블: scan_results, 컬럼에 entry_signal_strength + entry_signals JSON 포함)\n"
        "    │\n"
        "    ├─[6] HTML 리포트 생성                     → scanner/kr/reports/html_report.py:generate_daily_report\n"
        "    │       (data/reports/YYYY-MM-DD/index.html + ticker별 stock.html)\n"
        "    │\n"
        "    └─[7] 콘솔 요약 출력                       → scanner/cli.py (rich.Table)\n"
    )

    # ── 4. 모듈별 상세 ────────────────────────────────────────────
    _add_heading(doc, "4. 모듈별 상세 설명", 1)

    _add_heading(doc, "4.1 데이터 수집", 2)
    _add_table(
        doc,
        ["모듈", "역할", "핵심 함수"],
        [
            ["scanner/kr/fetcher.py",          "KIS OpenAPI 호출 — 일봉/1분봉/재무 데이터 수집", "fetch_daily_ohlcv, fetch_minute_ohlcv, fetch_fundamental"],
            ["scanner/kr/universe.py",         "KOSPI200 구성종목 갱신 (KIS 공개 종목마스터 zip)",  "update_kospi200"],
            ["scanner/data_pipeline.py",       "KR/US 통합 수집 dispatcher (시장별 fetcher 호출)", "run_data_pipeline, _fetch_ohlcv_one"],
        ],
        col_widths_cm=[5.5, 6.5, 4.5],
    )

    _add_heading(doc, "4.2 데이터 저장 (DB)", 2)
    _add_table(
        doc,
        ["모듈", "역할", "주요 모델/함수"],
        [
            ["scanner/db/models.py",     "SQLAlchemy 2.0 ORM 모델 정의",  "Universe, OHLCVDaily, OHLCVWeekly, OHLCVIntraday, ScanResult, Fundamental"],
            ["scanner/db/session.py",    "DB 세션 컨텍스트 매니저",       "get_session()"],
            ["scanner/db/repository.py", "스캔 결과 CRUD",                "save_scan_results, get_scan_results"],
            ["scanner/db/migrations.py", "DB 초기화 (테이블 + WAL)",      "init_database"],
            ["scanner/db/universe_db.py","활성 종목 헬퍼",                "get_active_tickers"],
        ],
        col_widths_cm=[5.0, 5.0, 6.5],
    )

    _add_heading(doc, "4.3 시간봉 합성 — 4시간봉 진입 타이밍 프레임", 2)
    _add_para(doc, "KIS 가 1분봉만 제공하므로 60분봉/4시간봉은 분석 시점에 합성한다.")
    _add_table(
        doc,
        ["함수", "설명"],
        [
            ["resample_to_minutes(df, rule, drop_partial)",   "일반화된 N분봉 합성. origin='start_day', offset='9h' → 자정이 아니라 09:00 영업시간 시작을 기점으로 그룹 경계를 끊음."],
            ["resample_to_4h(df, drop_partial=True)",         "1분봉 → 4시간봉. KOSPI 영업시간 기준 09:00~13:00 그룹 (완전봉) + 13:00~15:30 그룹 (부분봉, drop)."],
            ["resample_to_60min(df, drop_partial=True)",      "1분봉 → 60분봉. 정시 단위라 offset 영향 없음."],
        ],
        col_widths_cm=[7.0, 9.5],
    )
    _add_para(doc, "drop_partial 임계값 = 기대 분봉 수의 95%. KIS 1분봉이 09:01 부터 시작 (시초가 체결 직후) 하므로 09:00~12:59 그룹에 239봉 (240의 99.6%) 이 들어옴 — 이 자연 갭을 흡수한다.", italic=True)

    _add_heading(doc, "4.4 기술적 지표", 2)
    _add_table(
        doc,
        ["모듈", "지표", "용도"],
        [
            ["scanner/kr/indicators/moving_average.py", "SMA(N), EMA(N)",       "MA5/MA20/MA60 정배열 판정 + 골든크로스 탐지"],
            ["scanner/kr/indicators/rsi.py",            "RSI(14)",              "과매수/과매도 + 진입 신호의 RSI 반등 평가"],
            ["scanner/kr/indicators/macd.py",           "MACD(12, 26, 9)",      "히스토그램 양전환 (음→양) 평가"],
            ["scanner/kr/indicators/volume.py",         "volume_ratio",         "당일 거래량 / 직전 20일 평균 비율 (돌파 강도)"],
        ],
        col_widths_cm=[5.5, 4.0, 7.0],
    )

    _add_heading(doc, "4.5 패턴 탐지 — 핵심 모듈", 2)
    _add_table(
        doc,
        ["모듈", "역할"],
        [
            ["scanner/kr/patterns/base.py",   "PatternDetector ABC + EntrySignal/PatternResult dataclass + entry_signal() 추상"],
            ["scanner/kr/patterns/trend.py",  "주봉 추세 판정 (uptrend / downtrend / sideways) — detect_weekly_trend()"],
        ],
        col_widths_cm=[5.5, 11.0],
    )

    # ── 4.5.0 주봉 추세 판정 ──────────────────────────────────────
    _add_heading(doc, "4.5.0 주봉 추세 판정 (trend.py)", 3)
    _add_para(doc, "scanner/kr/patterns/trend.py 의 detect_weekly_trend() — 모든 패턴이 신뢰도 점수 산출 시 공유한다. 일봉을 W-MON 으로 resample 한 주봉을 입력으로 받아 3가지 신호로 방향을 판정한다.")
    _add_table(
        doc,
        ["판정 신호", "정의", "임계값"],
        [
            ["MA20 > MA60",        "단기-장기 이평선 정배열",                        "주봉 종가의 SMA20 / SMA60 비교"],
            ["종가 > MA20",        "가격이 단기선 위인지",                           "weekly close > weekly SMA20"],
            ["MA60 기울기",        "최근 8주 누적 변화율 (OLS slope × len / mean)",  "uptrend ≥ -1%, downtrend ≤ -2%"],
        ],
        col_widths_cm=[5.0, 6.5, 5.0],
    )
    _add_table(
        doc,
        ["방향", "조건"],
        [
            ["uptrend",          "정배열 AND 종가>MA20 AND MA60 8주 슬로프 ≥ -1%"],
            ["downtrend",        "역배열 AND 종가<MA20 AND MA60 8주 슬로프 ≤ -2%"],
            ["sideways",         "위 두 조건 모두 불충족"],
            ["insufficient_data","주봉 행 수 < 68 (= MA60 + 슬로프 윈도우 8) 또는 MA NaN"],
        ],
        col_widths_cm=[4.5, 12.0],
    )
    _add_para(doc, "추세 강도 (strength) = (MA20 − MA60) / MA60 × 100. 양수 = 정배열 폭, 음수 = 역배열 폭. 절댓값이 클수록 추세 진행도가 높다.", italic=True)

    # ── 4.5.1 쌍바닥 ──────────────────────────────────────────────
    _add_heading(doc, "4.5.1 쌍바닥 (DoubleBottomDetector — double_bottom.py)", 3)
    _add_para(doc, "U 자형 저점 두 개가 비슷한 가격에 형성되고, 사이의 고점(목선)을 돌파하면 매수 신호.")
    _add_para(doc, "데이터 요건: 일봉 ≥ 65 행 (60일 lookback + peak distance 5)", italic=True)
    _add_para(doc, "탐지 알고리즘:", bold=True)
    _add_code(doc,
        "1. 최근 60일 윈도우에서 종가의 저점 탐색\n"
        "    - scipy.signal.find_peaks(-close, distance=5, prominence=중앙값×3%)\n"
        "    - prominence: 중앙값 × 3% — 미세 노이즈 필터링\n"
        "    - distance=5: 저점 간 최소 5거래일 간격\n"
        "2. 최근 MAX_LOWS(=3) 개 저점만 검토 (트림)\n"
        "3. 저점 가격 균일성: (max − min) / min ≤ 6% (LOW_TOLERANCE × 2)\n"
        "4. 두 저점 사이 최고가 = 목선 (neckline)\n"
        "5. 목선 ≥ 평균 저점 × 1.05 (NECKLINE_GAIN_PCT = 5%)\n"
        "6. 두 번째 저점 이후 최근 5거래일 내 종가 > 목선 → 돌파 확인\n"
    )
    _add_para(doc, "진입가 / 손절가 / 목표가:", bold=True)
    _add_table(
        doc,
        ["가격", "산출식"],
        [
            ["진입가",  "neckline × 1.005 (목선 돌파 후 0.5% 위)"],
            ["손절가",  "평균 저점 × 0.97 (저점 -3%)"],
            ["목표가",  "neckline + (neckline − 평균 저점) — 패턴 높이만큼 가산"],
        ],
        col_widths_cm=[3.5, 13.0],
    )
    _add_para(doc, "raw_score 구성 (0~100):", bold=True)
    _add_table(
        doc,
        ["항목", "최대점", "산출"],
        [
            ["저점 균일성",         "30",  "1 − spread_pct / 6%"],
            ["목선 높이",           "30",  "min(1, (목선 − 평균저점) / 평균저점 / 10%)"],
            ["돌파 강도",           "20",  "min(1, (현재가 − 목선) / 목선 / 3%)"],
            ["돌파 후 거래량 비율", "20",  "min(1, (post_avg_vol / pre_avg_vol − 1) / 0.5)"],
            ["higher low 보너스",   "+5",  "두 번째 저점 ≥ 첫 저점 × 1.005 — 매도세 약화 신호"],
        ],
        col_widths_cm=[5.5, 2.0, 9.0],
    )

    # ── 4.5.2 골든크로스 ──────────────────────────────────────────
    _add_heading(doc, "4.5.2 골든크로스 (GoldenCrossDetector — golden_cross.py)", 3)
    _add_para(doc, "단기선이 장기선을 상향 돌파하는 추세 전환 신호. 단순 교차가 아니라 장기선의 평탄/상승 + 돌파일 거래량 급증을 함께 요구해 페이크아웃을 거른다.")
    _add_para(doc, "데이터 요건: 일봉 ≥ 60(MA60) + 5(RECENT) + 20(SLOPE) = 85 행", italic=True)
    _add_para(doc, "탐지 알고리즘:", bold=True)
    _add_code(doc,
        "1. SMA20, SMA60 계산\n"
        "2. 최근 5거래일 내 골든크로스 발생 인덱스 탐색\n"
        "    - prev: MA20[i-1] ≤ MA60[i-1]\n"
        "    - curr: MA20[i]   > MA60[i]\n"
        "3. MA60 기울기 검증 — 최근 20일의 (last − first) / first\n"
        "    - 슬로프 ≥ -0.5% 이어야 통과 (평탄/상승 인정)\n"
        "    - 하락 추세에서 일시 반등은 거름\n"
        "4. 돌파일 거래량 비율 ≥ 1.2 (당일 / 직전 20일 평균)\n"
    )
    _add_para(doc, "진입가 / 손절가 / 목표가:", bold=True)
    _add_table(
        doc,
        ["가격", "산출식"],
        [
            ["진입가",  "현재 종가"],
            ["손절가",  "MA60 × 0.97 (장기선 -3%)"],
            ["목표가",  "현재 종가 × 1.10 (+10%)"],
        ],
        col_widths_cm=[3.5, 13.0],
    )
    _add_para(doc, "raw_score 구성 (0~100):", bold=True)
    _add_table(
        doc,
        ["항목", "최대점", "산출"],
        [
            ["MA60 기울기",   "30", "min(1, max(0, (slope − -0.5%) / 2%))"],
            ["돌파일 거래량", "30", "min(1, (vol_ratio − 1) / 1)"],
            ["MA 이격",       "20", "min(1, (MA20 − MA60) / MA60 / 3%)"],
            ["교차 신선도",   "20", "max(0, 1 − days_since_cross / 5)"],
        ],
        col_widths_cm=[5.5, 2.0, 9.0],
    )

    # ── 4.5.3 박스권 돌파 ─────────────────────────────────────────
    _add_heading(doc, "4.5.3 박스권 돌파 (BoxBreakoutDetector — box_breakout.py)", 3)
    _add_para(doc, "수평 횡보 후 상단 돌파. 30/45/60일 윈도우 중 가장 명확한(range_pct 최소) 박스를 자동 선택한다.")
    _add_para(doc, "데이터 요건: 일봉 ≥ 60 + 3 + 1 = 64 행", italic=True)
    _add_para(doc, "탐지 알고리즘:", bold=True)
    _add_code(doc,
        "1. 윈도우 후보 [30, 45, 60] 각각 시도, range_pct 최소인 박스 채택\n"
        "2. 박스 영역 = (전체 길이 − RECENT_DAYS=3 − 1) 직전의 window 일\n"
        "    - box_top    = 영역 내 close 최댓값  (※ 종가 기반 — wick outlier 제거)\n"
        "    - box_bottom = 영역 내 close 최솟값\n"
        "    - range_pct  = (top − bottom) / 박스 평균 종가\n"
        "    - range_pct ≤ 15% 이어야 박스 인정 (BOX_BREAKOUT_RANGE_PCT, KOSPI 현실 반영)\n"
        "3. 최근 3일 내 종가 > box_top × 1.01 → 돌파 확인\n"
        "4. 돌파일 거래량 비율 ≥ 1.5 (당일 / 직전 20일 평균)\n"
    )
    _add_para(doc, "주의 — high.max/low.min 기반 산출은 단일 wick 영향으로 박스가 부풀려져 KOSPI200 의 30일 range median 24% 를 만들어 박스 후보 0건 이슈를 일으킴. 종가 기반 산출 + 15% 임계값 조합이 KOSPI 종목의 자연스런 변동폭과 박스의 의미(매수/매도 종가 균형) 를 모두 반영. 진단: scripts/diag_box_breakout.py / diag_box_breakout_v2.py.", italic=True)
    _add_para(doc, "진입가 / 손절가 / 목표가:", bold=True)
    _add_table(
        doc,
        ["가격", "산출식"],
        [
            ["진입가",  "현재 종가"],
            ["손절가",  "box_bottom × 0.98 (박스 하단 -2%)"],
            ["목표가",  "box_top + (box_top − box_bottom) — 박스 높이만큼 추가"],
        ],
        col_widths_cm=[3.5, 13.0],
    )
    _add_para(doc, "raw_score 구성 (0~100):", bold=True)
    _add_table(
        doc,
        ["항목", "최대점", "산출"],
        [
            ["박스 밀도",     "30", "1 − range_pct / 10%"],
            ["돌파일 거래량", "30", "min(1, (vol_ratio − 1) / 1)"],
            ["돌파 강도",     "20", "min(1, (현재가 − box_top) / box_top / 5%)"],
            ["돌파 신선도",   "20", "max(0, 1 − days_since_break / 3)"],
        ],
        col_widths_cm=[5.5, 2.0, 9.0],
    )

    # ── 4.5.4 눌림목 ──────────────────────────────────────────────
    _add_heading(doc, "4.5.4 눌림목 (PullbackDetector — pullback.py)", 3)
    _add_para(doc, "상승 추세 종목이 일시 조정으로 이평선까지 눌렸다가 양봉 + 거래량으로 반등하는 자리. 4가지 패턴 중 유일하게 주봉 추세가 상승이어야 한다는 필수 전제 조건이 detect 단계에 들어간다.")
    _add_para(doc, "데이터 요건: 일봉 ≥ 350 행 (주봉 추세 64주 × 5일 + 여유)", italic=True)
    _add_para(doc, "탐지 알고리즘:", bold=True)
    _add_code(doc,
        "1. 일봉 → 주봉 resample (W-MON, 미완성 주봉 drop)\n"
        "2. detect_weekly_trend() == 'uptrend' (필수 — 다른 패턴엔 없는 강제 조건)\n"
        "3. 일봉 정배열 확인: MA5 > MA20 > MA60\n"
        "4. 현재가가 MA20 또는 MA60 의 ±2% 범위 안에 있는지\n"
        "    - near_ma20 = |close − MA20| / MA20 ≤ 2%\n"
        "    - near_ma60 = |close − MA60| / MA60 ≤ 2%\n"
        "5. 당일 양봉 (close > open)\n"
        "6. 당일 거래량 > 직전 5일 평균\n"
    )
    _add_para(doc, "진입가 / 손절가 / 목표가:", bold=True)
    _add_table(
        doc,
        ["가격", "산출식"],
        [
            ["진입가",  "현재 종가"],
            ["손절가",  "MA60 × 0.97 (장기선 -3%)"],
            ["목표가",  "직전 20일 고점 (없거나 < 진입가면 진입가 × 1.08)"],
        ],
        col_widths_cm=[3.5, 13.0],
    )
    _add_para(doc, "raw_score 구성 (0~100):", bold=True)
    _add_table(
        doc,
        ["항목", "최대점", "산출"],
        [
            ["주봉 강도",     "30", "min(1, max(0, weekly_strength / 5%))"],
            ["MA 정배열 간격", "25", "min(1, (MA5−MA20) / MA20 / 2%)"],
            ["MA 근접도",     "25", "max(0, 1 − proximity_pct / 2%) — 가까울수록 고점수"],
            ["거래량",        "20", "min(1, (vol_ratio − 1) / 0.5)"],
        ],
        col_widths_cm=[5.5, 2.0, 9.0],
    )

    # ── 4.5.5 패턴 임계값 일람 ────────────────────────────────────
    _add_heading(doc, "4.5.5 패턴 임계값 일람 (config.py)", 3)
    _add_para(doc, "scanner/config.py 에 모든 패턴 임계값이 상수로 모여 있다 (매직 넘버 제거). 변경 시 한 곳만 수정.")
    _add_table(
        doc,
        ["상수", "값", "의미"],
        [
            ["DOUBLE_BOTTOM_LOOKBACK_DAYS",       "60",       "쌍바닥 탐색 구간"],
            ["DOUBLE_BOTTOM_LOW_TOLERANCE_PCT",   "0.03",     "두 저점 허용 차이 ±3%"],
            ["DOUBLE_BOTTOM_NECKLINE_GAIN_PCT",   "0.05",     "목선 +5% 이상"],
            ["GOLDEN_CROSS_FAST_MA / SLOW_MA",    "20 / 60",  "단기 / 장기 이평선"],
            ["GOLDEN_CROSS_RECENT_DAYS",          "5",        "최근 5일 내 크로스"],
            ["GOLDEN_CROSS_VOLUME_RATIO",         "1.2",      "돌파일 거래량 / 20일 평균"],
            ["BOX_BREAKOUT_MIN/MAX_DAYS",         "30 / 60",  "박스 형성 일수 범위"],
            ["BOX_BREAKOUT_RANGE_PCT",            "0.15",     "박스 폭 ±15% (종가 기반, KOSPI 현실 반영)"],
            ["BOX_BREAKOUT_BREAK_PCT",            "0.01",     "상단 돌파 +1%"],
            ["BOX_BREAKOUT_VOLUME_RATIO",         "1.5",      "돌파일 거래량 / 20일 평균"],
            ["PULLBACK_MA_NEAR_PCT",              "0.02",     "MA20/MA60 ±2% 근접"],
            ["PULLBACK_VOLUME_LOOKBACK",          "5",        "거래량 평균 5일"],
            ["MA_SHORT / MEDIUM / LONG",          "5 / 20 / 60", "정배열 검사 이평선"],
        ],
        col_widths_cm=[6.0, 2.5, 8.0],
    )

    _add_heading(doc, "4.6 신뢰도 점수 (CLAUDE.md §6)", 2)
    _add_para(doc, "scanner/kr/scoring/scorer.py 의 calculate_confidence_score() — 5개 컴포넌트의 0~100 점수를 가중합산하여 최종 신뢰도 점수를 산출.")
    _add_table(
        doc,
        ["컴포넌트", "가중치", "0~100 점수 산출"],
        [
            ["주봉 추세 (weekly_trend)",       "30%", "uptrend=100 / sideways=40 / downtrend=0 / 기타=0"],
            ["패턴 명확도 (pattern_clarity)",  "25%", "PatternResult.raw_score 그대로 사용 (각 detector 의 _calc_raw_score)"],
            ["거래량 (volume)",                "20%", "min(100, max(0, (vol_ratio − 1) / 0.5 × 100)) — 1.0=0점, 1.5+=100점"],
            ["MA 정배열 (ma_alignment)",       "15%", "details.ma5/20/60 정배열이면 100, 아니면 0. 없으면 100 (탐지 자체가 정배열 함의)"],
            ["RSI (rsi)",                      "10%", "40~70=100점, <40=val/40×100, >70=(100−val)/30×100, 데이터 부족=중립 50"],
            ["**합계**",                       "100%", "Σ(score_i × weight_i), clip [0, 100], 소수 둘째 자리 round"],
        ],
        col_widths_cm=[5.5, 2.0, 9.0],
    )
    _add_para(doc, "주의 — 거래량 점수의 입력 vol_ratio 는 detector.details['vol_ratio'] 에서 가져온다. 패턴별로 의미가 약간 다름:", italic=True)
    _add_table(
        doc,
        ["패턴", "vol_ratio 의미"],
        [
            ["double_bottom",  "details 에 vol_ratio 키 없음 → scorer 가 0 으로 받아 거래량 점수 0 (raw_score 의 거래량 항목으로 대체 평가)"],
            ["golden_cross",   "details.vol_ratio_at_cross — 돌파일 거래량 / 20일 평균"],
            ["box_breakout",   "details.vol_ratio_at_break — 동일"],
            ["pullback",       "details.vol_ratio — 당일 거래량 / 5일 평균"],
        ],
        col_widths_cm=[4.0, 12.5],
    )
    _add_para(doc, "70점 이상 종목만 최종 리포트에 포함한다 (CLI --min-confidence 로 조정).", italic=True)

    _add_heading(doc, "4.7 진입 신호 (4시간봉 기반)", 2)
    _add_para(doc, "각 detector 의 entry_signal() 메서드가 4가지 신호를 평가한다 (각 25점, 합계 0~100). 4시간봉이 진입 타이밍 프레임이며, MA20·박스 상단·20일 고점 같은 일봉 컨텍스트는 일봉 그대로 사용한다.")
    _add_table(
        doc,
        ["패턴", "1. RSI", "2. 거래량 (4h)", "3. MACD (일봉)", "4. 패턴별 신호"],
        [
            ["double_bottom",  "rsi_bounce: 50 미만에서 상승",  "bullish_volume",  "macd_cross",     "prev_high_break (20일 고점 돌파)"],
            ["golden_cross",   "rsi_rising: 30~70 상승",        "bullish_volume",  "macd_positive",  "above_ma20"],
            ["box_breakout",   "rsi_above_50: 50 상향 돌파",    "bullish_volume",  "macd_positive",  "above_box_top (박스 상단 위)"],
            ["pullback",       "rsi_bounce: 40~70 상승",        "bullish_volume",  "macd_positive",  "above_ma20"],
        ],
        col_widths_cm=[3.0, 3.5, 3.0, 3.0, 4.0],
    )

    _add_heading(doc, "4.8 필터", 2)
    _add_table(
        doc,
        ["모듈", "필터", "조건"],
        [
            ["scanner/kr/filtering/volume_filter.py",      "거래량 모멘텀",   "최근 5일 거래량 평균 > 직전 20일 거래량 평균"],
            ["scanner/kr/filtering/fundamental_filter.py", "시가총액",       "≥ 1,000억 원 (KOSPI200 우량주 기준)"],
        ],
        col_widths_cm=[6.5, 3.5, 6.5],
    )
    _add_para(doc, "PER / 부채비율은 분석에 사용하지 않는다 (데이터 수집은 유지). 거래대금 컷은 KOSPI200 단계에서 사실상 모두 통과해 비활성화 — KOSPI/KOSDAQ 일반 종목 도입 시 재활성화 검토.", italic=True)

    _add_heading(doc, "4.9 호가단위 정합", 2)
    _add_table(
        doc,
        ["모듈", "역할"],
        [
            ["scanner/kr/quote_align.py", "한국거래소 가격대별 호가단위 (1원/5원/10원/50원/100원/500원/1000원) 로 진입가/손절가/목표가를 정합 — KIS 주문 가능한 가격으로 정렬."],
        ],
        col_widths_cm=[5.0, 11.5],
    )

    _add_heading(doc, "4.10 리포트 생성", 2)
    _add_table(
        doc,
        ["모듈", "출력 형식", "내용"],
        [
            ["scanner/kr/reports/html_report.py",       "HTML",      "data/reports/YYYY-MM-DD/index.html + ticker별 stock.html (LightweightCharts v5)"],
            ["scanner/kr/reports/excel_export.py",      "CSV / XLSX","data/exports/YYYY-MM-DD.{csv,xlsx} — 종목/패턴/점수/진입가/손절가/목표가"],
            ["scanner/kr/reports/comment_generator.py", "한국어 텍스트","패턴별 자동 코멘트 (AI 생성 X — 규칙 기반 템플릿)"],
        ],
        col_widths_cm=[6.0, 3.0, 7.5],
    )

    _add_heading(doc, "4.11 백테스트", 2)
    _add_table(
        doc,
        ["모듈", "역할", "핵심 함수"],
        [
            ["scanner/kr/backtest/engine.py", "DB 의 과거 ScanResult 를 재생해 hold_days 안에 목표가/손절가 도달 여부를 시뮬레이션. 승률/평균 수익률/수익 팩터/최대 낙폭을 산출.", "run_backtest(pattern_name, period_days, hold_days, min_score)"],
        ],
        col_widths_cm=[5.5, 7.5, 3.5],
    )

    _add_heading(doc, "4.12 진입점 (CLI / API)", 2)
    _add_table(
        doc,
        ["모듈", "역할"],
        [
            ["scanner/cli.py",                  "typer CLI — scan / results / show / report / export / backtest / serve / fetch / fetch-all / update-universe"],
            ["scanner/pipeline.py",             "run_daily_pipeline() — CLI 와 별개로 직접 호출 가능한 일일 스캔 함수"],
            ["scanner/api/main.py",             "FastAPI app — uvicorn serve 진입점"],
            ["scanner/api/routers/stocks.py",   "/api/stocks/{ticker}/ohlcv (일봉+주봉+4시간봉) + /api/stocks/{ticker}/analysis (스캔 결과 + 진입 신호)"],
        ],
        col_widths_cm=[5.5, 11.0],
    )

    # ── 5. 데이터 모델 ────────────────────────────────────────────
    _add_heading(doc, "5. 핵심 DB 테이블 (scanner/db/models.py)", 1)
    _add_table(
        doc,
        ["테이블", "주요 컬럼", "설명"],
        [
            ["universe",         "ticker, name, market, market_cap, is_active, updated_at",                                "활성 종목 목록 (KOSPI200)"],
            ["ohlcv_daily",      "ticker, date, open/high/low/close, volume, value",                                       "일봉 OHLCV (+ 거래대금)"],
            ["ohlcv_weekly",     "ticker, week_start_date, OHLC, volume",                                                  "주봉 (일봉 resample 결과)"],
            ["ohlcv_intraday",   "ticker, datetime, OHLC, volume",                                                         "1분봉 (4시간봉 합성 원천)"],
            ["fundamentals",     "ticker, date, per, debt_ratio",                                                          "재무 (현재 분석엔 미사용, 데이터만 적재)"],
            ["scan_results",     "scan_date, ticker, pattern_name, confidence_score, entry/stop/target, entry_signal_strength, entry_signals(JSON), passed_filters", "일일 스캔 결과 — 점수 + 진입 가격 + 진입 신호"],
            ["backtest_results", "pattern_name, period, win_rate, avg_return_pct, ...",                                    "백테스트 시뮬레이션 결과"],
        ],
        col_widths_cm=[3.0, 7.5, 6.0],
    )

    # ── 6. 모듈 인덱스 ────────────────────────────────────────────
    _add_heading(doc, "6. 모듈 빠른 인덱스 (알파벳 순)", 1)
    _add_table(
        doc,
        ["경로", "한 줄 요약"],
        [
            ["scanner/api/main.py",                              "FastAPI 앱"],
            ["scanner/api/routers/stocks.py",                    "/api/stocks/{ticker}/{ohlcv,analysis}"],
            ["scanner/cli.py",                                   "typer CLI 진입점 (scan/results/show/...)"],
            ["scanner/config.py",                                "임계값 상수 + 환경변수 로딩"],
            ["scanner/data_pipeline.py",                         "통합 데이터 수집 dispatcher"],
            ["scanner/db/migrations.py",                         "DB 초기화 + 테이블 생성"],
            ["scanner/db/models.py",                             "SQLAlchemy ORM 모델"],
            ["scanner/db/repository.py",                         "ScanResult CRUD"],
            ["scanner/db/session.py",                            "get_session 컨텍스트 매니저"],
            ["scanner/db/universe_db.py",                        "활성 종목 조회 헬퍼"],
            ["scanner/kr/backtest/engine.py",                    "백테스트 엔진"],
            ["scanner/kr/fetcher.py",                            "KIS OpenAPI fetcher"],
            ["scanner/kr/filtering/fundamental_filter.py",       "시가총액 필터"],
            ["scanner/kr/filtering/volume_filter.py",            "거래량 모멘텀 필터"],
            ["scanner/kr/indicators/macd.py",                    "MACD"],
            ["scanner/kr/indicators/moving_average.py",          "SMA / EMA"],
            ["scanner/kr/indicators/rsi.py",                     "RSI"],
            ["scanner/kr/indicators/volume.py",                  "거래량 비율"],
            ["scanner/kr/intraday.py",                           "1분봉 → 60min/4h 합성"],
            ["scanner/kr/patterns/base.py",                      "패턴 추상 + EntrySignal/PatternResult"],
            ["scanner/kr/patterns/box_breakout.py",              "박스권 돌파"],
            ["scanner/kr/patterns/double_bottom.py",             "쌍바닥"],
            ["scanner/kr/patterns/golden_cross.py",              "골든크로스"],
            ["scanner/kr/patterns/pullback.py",                  "눌림목"],
            ["scanner/kr/patterns/trend.py",                     "주봉 추세 판정"],
            ["scanner/kr/quote_align.py",                        "호가단위 정합"],
            ["scanner/kr/reports/comment_generator.py",          "패턴별 한국어 코멘트"],
            ["scanner/kr/reports/excel_export.py",               "CSV/XLSX 내보내기"],
            ["scanner/kr/reports/html_report.py",                "HTML 리포트 (Jinja2)"],
            ["scanner/kr/scanner.py",                            "analyze_ticker / scan_universe"],
            ["scanner/kr/scoring/scorer.py",                     "신뢰도 점수 계산"],
            ["scanner/kr/universe.py",                           "KOSPI200 갱신"],
            ["scanner/pipeline.py",                              "일일 스캔 오케스트레이터"],
        ],
        col_widths_cm=[7.0, 9.5],
    )

    # ── 7. 부록 ────────────────────────────────────────────────────
    _add_heading(doc, "7. 부록 — Phase C-3b 4시간봉 전환 요약", 1)
    _add_para(doc, "본 브랜치 (feat/web-entry-signal-and-4h-toggle) 의 변경 핵심:")
    _add_table(
        doc,
        ["변경 영역",         "변경 내용"],
        [
            ["scanner/kr/intraday.py",                "4h 그룹 경계를 09:00 영업시간 시작 기준으로 끊기 (origin='start_day', offset='9h'). drop_partial 임계값 95% (KIS 1분봉이 09:01 시작이라 239봉만 채워지는 자연 갭 흡수)."],
            ["scanner/pipeline.py",                   "_load_intraday_dfs() 가 resample_to_60min → resample_to_4h 로 전환. 일별 1봉 (09:00~12:59 완전봉) 만 사용."],
            ["scanner/cli.py",                        "scan() 명령에 _load_intraday_dfs() 호출 + analyze_ticker 에 intraday_df 전달 흐름 추가 (이전엔 누락)."],
            ["scanner/kr/patterns/*.py (4개)",        "entry_signal() 의 bullish_volume 평가를 4시간봉 기반으로 변경 (이전엔 일봉). RSI 는 이미 intraday 기반. MACD·MA20·박스 상단 등 일봉 컨텍스트는 유지."],
            ["scanner/kr/scanner.py / patterns/base.py", "docstring 의 '60분봉' 표기를 '4시간봉' 으로 정정 (CLAUDE.md §1 정합성)."],
            ["scanner/api/routers/stocks.py",         "/api/stocks/{ticker}/ohlcv 응답에 4h 봉 추가 (KR 종목 + 분봉 적재 시)."],
            ["scanner/kr/reports/html_report.py",     "_serialize_time 이 datetime 을 UNIX 타임스탬프 + KST(+9h) 시프트로 직렬화 — LightweightCharts 가 UTC 표시값을 그대로 KST 시각으로 보여주는 트릭."],
            ["web/stock.html",                        "[4H] 토글 버튼 추가, 진입 강도 + 진입 신호 chip 표시, _formatTime 이 number 타임스탬프를 MM-DD HH:00 형식으로 출력."],
        ],
        col_widths_cm=[5.5, 11.0],
    )

    # ── 저장 ──────────────────────────────────────────────────────
    out = Path(__file__).resolve().parent.parent / "KR_모듈_흐름.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"저장 완료: {path}")
